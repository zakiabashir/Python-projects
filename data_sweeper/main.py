import streamlit as st
import pandas as pd
from docx import Document
import fitz  # PyMuPDF for PDF
from PIL import Image  # For images
import io

def convert_file(file, output_format):
    try:
        # Read the input file
        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.name.endswith('.xlsx') or file.name.endswith('.xls'):
            df = pd.read_excel(file)
        elif file.name.endswith('.json'):
            df = pd.read_json(file)
        elif file.name.endswith('.docx'):
            doc = Document(file)
            text = [paragraph.text for paragraph in doc.paragraphs]
            df = pd.DataFrame(text, columns=['Content'])
        elif file.name.endswith('.pdf'):
            pdf = fitz.open(stream=file.read(), filetype="pdf")
            text = ""
            for page in pdf:
                text += page.get_text()
            df = pd.DataFrame([text], columns=['Content'])
        elif file.name.endswith(('.jpg', '.jpeg', '.png')):
            img = Image.open(file)
            # Convert image to bytes
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format=img.format)
            img_byte_arr = img_byte_arr.getvalue()
            df = pd.DataFrame([{'Image': img_byte_arr}])
        else:
            st.error("Unsupported file format")
            return None
            
        # Convert to selected format
        if output_format == 'CSV':
            output = df.to_csv(index=False)
            mime = 'text/csv'
            file_extension = '.csv'
        elif output_format == 'Excel':
            output = df.to_excel(index=False)
            mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            file_extension = '.xlsx'
        elif output_format == 'JSON':
            output = df.to_json()
            mime = 'application/json'
            file_extension = '.json'
        elif output_format == 'DOCX':
            doc = Document()
            for _, row in df.iterrows():
                doc.add_paragraph(str(row[0]))
            doc_io = io.BytesIO()
            doc.save(doc_io)
            output = doc_io.getvalue()
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_extension = '.docx'
        elif output_format == 'PDF':
            pdf = fitz.open()
            page = pdf.new_page()
            for _, row in df.iterrows():
                page.insert_text((50, 50), str(row[0]))
            pdf_bytes = pdf.write()
            output = pdf_bytes
            mime = 'application/pdf'
            file_extension = '.pdf'
        elif output_format in ['JPG', 'PNG']:
            if 'Image' in df.columns:
                output = df['Image'].iloc[0]
                mime = 'image/jpeg' if output_format == 'JPG' else 'image/png'
                file_extension = '.jpg' if output_format == 'JPG' else '.png'
            else:
                st.error("Cannot convert to image format")
                return None
            
        return output, mime, file_extension
        
    except Exception as e:
        st.error(f"Error converting file: {str(e)}")
        return None

def main():
    st.title('File Converter App')
    st.write('Convert your files between different formats')
    
    # File upload
    uploaded_file = st.file_uploader("Choose a file", type=['csv', 'xlsx', 'xls', 'json', 'docx', 'pdf', 'jpg', 'jpeg', 'png'])
    
    if uploaded_file is not None:
        # Show original file details
        st.subheader('Original File Details')
        st.write(f"Filename: {uploaded_file.name}")
        
        # Select output format
        output_format = st.selectbox(
            'Select output format',
            ['CSV', 'Excel', 'JSON', 'DOCX', 'PDF', 'JPG', 'PNG']
        )
        
        # Convert button
        if st.button('Convert File'):
            result = convert_file(uploaded_file, output_format)
            
            if result:
                output, mime, file_extension = result
                
                # Download converted file
                st.download_button(
                    label=f"Download converted file as {output_format}",
                    data=output,
                    file_name=f"converted{file_extension}",
                    mime=mime
                )
                
                st.success('File converted successfully!')

if __name__ == '__main__':
    main()
